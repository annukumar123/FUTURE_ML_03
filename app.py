import os
import spacy
from flask import Flask, request, render_template
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import fitz  # PyMuPDF for better PDF handling
from docx import Document
from sentence_transformers import SentenceTransformer, util
from docx import Document
import io
from flask import send_file, session

# Load the model (do this outside the route so it only loads once)
model = SentenceTransformer('all-MiniLM-L6-v2')


app = Flask(__name__)
app.secret_key = 'resume_screener_secret'
app.config['UPLOAD_FOLDER'] = 'uploads/'
nlp = spacy.load("en_core_web_md")

# Ensure upload folder exists
if not os.path.exists(app.config['UPLOAD_FOLDER']):
    os.makedirs(app.config['UPLOAD_FOLDER'])

def extract_text(file_path):
    ext = file_path.split('.')[-1].lower()
    text = ""
    if ext == 'pdf':
        with fitz.open(file_path) as doc:
            for page in doc: text += page.get_text()
    elif ext == 'docx':
        doc = Document(file_path)
        text = "\n".join([p.text for p in doc.paragraphs])
    elif ext == 'txt':
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
    return text

def get_suggestions(resume_text, jd_text):
    
    jd_doc = nlp(jd_text.lower())
    resume_doc = nlp(resume_text.lower())
    
    # 1. Define words to ignore (Non-technical noise)
    exclude_words = {
        'responsibilities', 'commands', 'end', 'degree', 'page', 'date', 
        'experience', 'knowledge', 'ability', 'role', 'work', 'team', 
        'candidate', 'requirements', 'details', 'eligibility', 'others',
        'graduate', 'exposure', 'projects', 'field', 'mindset', 'commitment',
        'discipline', 'curiosity', 'baseline', 'growth', 'opportunity', 'scenarios'
    }
    
    # 2. Extract nouns/proper nouns as "keywords"
    # Added a check to ensure the word isn't a stopword AND isn't in our exclude list
    jd_keywords = set([
        t.text for t in jd_doc 
        if t.pos_ in ['NOUN', 'PROPN'] 
        and not t.is_stop 
        and len(t.text) > 2
        and t.text not in exclude_words
    ])
    
    res_keywords = set([
        t.text for t in resume_doc 
        if t.pos_ in ['NOUN', 'PROPN']
    ])
    
    # 3. Calculate the gap
    missing = list(jd_keywords - res_keywords)
    
    # Return the top 8 unique missing words
    return missing[:8]

@app.route('/')
def index():
    return render_template('matchresume.html')




from docx import Document
import io
from flask import send_file, session

# 1. Function to create the modified document
def generate_modified_docx(original_text, missing_keywords):
    doc = Document()
    
    # Add original content (Simplified for demonstration)
    doc.add_heading('Resume (Optimized)', 0)
    doc.add_paragraph(original_text)
    
    # Add the Optimization Section
    doc.add_page_break()
    doc.add_heading('ATS Optimization Section', level=1)
    doc.add_paragraph("The following keywords have been added to align with the Job Description requirements:")
    
    # Add keywords as bold text
    p = doc.add_paragraph()
    p.add_run(", ".join(missing_keywords)).bold = True
    
    # Save to memory
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

# 2. Route to handle the download
@app.route('/download_optimized/<int:index>')
def download_optimized(index):
    # Retrieve data from session (stored during /matcher)
    results = session.get('last_results', [])
    if index < len(results):
        res = results[index]
        modified_file = generate_modified_docx(res['full_text'], res['suggestions'])
        
        return send_file(
            modified_file, 
            as_attachment=True, 
            download_name=f"Optimized_{res['filename']}.docx"
        )
    return "Error: File data not found.", 404



@app.route('/matcher', methods=['POST'])
def matcher():
    job_description = request.form.get('job_description')
    resume_files = request.files.getlist('resumes')
    
    results = []
    resume_texts = []
    filenames = []

    for file in resume_files:
        path = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
        file.save(path)
        text = extract_text(path)
        resume_texts.append(text)
        filenames.append(file.filename)

    if not resume_texts or not job_description:
        return "Please upload resumes and provide a JD."

    # Scoring Logic
    vectorizer = TfidfVectorizer()
    all_content = [job_description] + resume_texts
    # tfidf_matrix = vectorizer.fit_transform(all_content)
    embeddings = model.encode([job_description] + resume_texts)
    
    # Compare JD (index 0) against all Resumes (index 1 onwards)
    cosine_scores = util.cos_sim(embeddings[0], embeddings[1:])[0]
    # similarities = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:])[0]

    for i in range(len(filenames)):
        score = round(float(cosine_scores[i]) * 100, 2)
        missing_skills = get_suggestions(resume_texts[i], job_description)
        
        results.append({
            'filename': filenames[i],
            'score': score,
            'suggestions': missing_skills,
            'full_text': resume_texts[i]  # Store full text for later use in optimized doc
        })

    # Sort by score
    results = sorted(results, key=lambda x: x['score'], reverse=True)
    session['last_results'] = results
    # Optional: Clean up files after processing
    for file in resume_files:
         os.remove(os.path.join(app.config['UPLOAD_FOLDER'], file.filename))
    return render_template('results.html', results=results)

if __name__ == '__main__':
    app.run(debug=True)